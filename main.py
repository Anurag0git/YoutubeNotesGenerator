from flask import Flask, render_template, request, jsonify
import google.generativeai as genai
from youtube_transcript_api import YouTubeTranscriptApi, TranscriptsDisabled, NoTranscriptFound
import re
import requests
from docx import Document
import os
from dotenv import load_dotenv

load_dotenv()

app = Flask(__name__)

# Set up Gemini API
GEMINI_API_KEY = os.getenv("GOOGLE_API_KEY")
if not GEMINI_API_KEY:
    raise ValueError("GEMINI_API_KEY not set in environment variables.")

genai.configure(api_key=GEMINI_API_KEY)


# Function to extract video ID
def extract_video_id(youtube_url):
    pattern = r"(?:https?:\/\/)?(?:www\.)?(?:youtube\.com\/(?:[^\/]+\/.+\/|(?:v|e(?:mbed)?)\/|.*[?&]v=)|youtu\.be\/)([^\"&?\/\s]{11})"
    match = re.search(pattern, youtube_url)
    return match.group(1) if match else None


# API Route to Fetch Transcript with Better Error Handling
@app.route('/get_transcript', methods=['POST'])
def get_transcript():
    data = request.json
    video_url = data.get("url")
    video_id = extract_video_id(video_url)

    if not video_id:
        return jsonify({"error": "Invalid YouTube URL"}), 400

    try:
        transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)

        # Prioritized languages
        preferred_languages = ["en", "en-IN", "hi", "hi-IN"]
        selected_transcript = None

        # Try to fetch transcript in preferred order
        for lang in preferred_languages:
            try:
                selected_transcript = transcript_list.find_transcript([lang])
                break
            except:
                continue

        # If no preferred language is found, get any available transcript
        if not selected_transcript:
            selected_transcript = transcript_list.find_generated_transcript(
                [t.language_code for t in transcript_list]
            )

        # Extract transcript text properly
        text = " ".join([entry.text for entry in selected_transcript.fetch()])

        return jsonify({"transcript": text, "language": selected_transcript.language})

    except NoTranscriptFound:
        return jsonify({"error": "No transcript available for this video"}), 400
    except TranscriptsDisabled:
        return jsonify({"error": "Transcripts are disabled for this video"}), 400
    except requests.exceptions.RequestException as e:
        return jsonify({"error": f"YouTube request failed: {str(e)}"}), 500
    except Exception as e:
        return jsonify({"error": f"Could not fetch transcript: {str(e)}"}), 500

# API Route to Generate Summary
@app.route('/generate_summary', methods=['POST'])
def generate_summary():
    data = request.json
    transcript = data.get("transcript", "").strip()

    if not transcript:
        return jsonify({"error": "Transcript is empty"}), 400

    prompt = f"Generate detailed notes for each section in the given transcript in simplest language English and do not use formatting, use easiest explaination so that even kid could understand, make proper notes so once read, I should not forget them in my life:\n\n{transcript}"

    try:
        model = genai.GenerativeModel("gemini-2.0-flash")
        response = model.generate_content(prompt)
        summary = response.text if response else "Failed to generate summary."
        return jsonify({"summary": summary})
    except Exception as e:
        return jsonify({"error": f"Could not generate summary: {str(e)}"}), 500


# API Route to Save Summary as .docx
@app.route('/save_summary', methods=['POST'])
def save_summary():
    data = request.json
    summary = data.get("summary", "").strip()

    if not summary:
        return jsonify({"error": "No summary to save!"}), 400

    file_path = "summary.docx"
    doc = Document()
    doc.add_heading("YouTube Video Summary", level=1)
    doc.add_paragraph(summary)
    doc.save(file_path)

    return jsonify({"message": "Summary saved successfully!", "file": file_path})


# Serve the Webpage
@app.route('/')
def home():
    return render_template('index.html')



if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5555))
    app.run(host="0.0.0.0", port=port, debug=True)
